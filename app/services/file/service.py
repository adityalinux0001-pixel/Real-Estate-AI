from pathlib import Path
from fastapi import HTTPException
from app.services.ai.rag import clear_vectors
from app.models.file import File
from sqlmodel.ext.asyncio.session import AsyncSession
import docx
import asyncio
import logging

logger = logging.getLogger(__name__)


async def clear_file(file_id: int, path: str) -> bool:
    try:
        success = await clear_vectors(file_id)
        if not success:
            raise HTTPException(status_code=500, detail="Failed to clear old vectors from Pinecone")

        old_path = Path(path)
        if old_path.exists():
            await asyncio.to_thread(old_path.unlink)
            
        logger.info(f"Clearing file and vectors for file_id {file_id} returned: {success}")
        return True
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error deleting file: {e}")


async def save_file_to_db(name: str, path: str, category: str, building_id: int | None, user_id: int, db: AsyncSession) -> File:
    db_file = File(
        filename=name,
        path=str(path),
        category=category,
        building_id=building_id,
        user_id=user_id
    )
    db.add(db_file)
    await db.commit()
    await db.refresh(db_file)
    return db_file


async def save_text_as_docx(text: str, output_path: str, key: str) -> None:
    doc = docx.Document()

    if key == "structured":
        doc.add_heading("Structured Document", level=1)
    if key == "lease_abstract":
        doc.add_heading("Commercial Office Lease Abstract", level=1)
    if key == "lease_content":
        doc.add_heading("Legal Lease Document", level=1)
    if key == "report_summary":
        doc.add_heading("Complete Report Summary", level=1)
        
    for line in text.splitlines():
        line = line.strip()
        if not line:
            continue
        if line.startswith("•") and key == "structured":
            doc.add_paragraph(line.lstrip("• ").strip(), style="List Bullet")
        elif line.endswith(":") and key == "structured":
            doc.add_heading(line, level=2)
        elif line.startswith(("I.", "II.", "III.")) and key == "lease_abstract":
            doc.add_heading(line, level=2)
        else:
            doc.add_paragraph(line)
    doc.save(output_path)
    
