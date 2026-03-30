import asyncio
from pathlib import Path
import PyPDF2
from app.core.config import get_settings
import docx
import pandas as pd
import aiofiles

settings = get_settings()

async def extract_docx_text(docx_path: Path) -> str:
    doc = docx.Document(docx_path)
    full_text = []
    current_heading = None
    doc_element = doc.element.body
    
    for elem in doc_element:
        if elem.tag.endswith('p'):  
            para = doc.paragraphs[[p._element for p in doc.paragraphs].index(elem)]
            text = para.text.strip()
            if not text:
                continue
                
            if para.style.name.startswith('Heading') or (para.runs and any(run.bold for run in para.runs)):
                if current_heading and full_text and full_text[-1] != '':
                    full_text.append('')
                current_heading = text
                full_text.append(text + ':')
            else:
                if text:
                    full_text.append(text)
                    if full_text[-1] != '':
                        full_text.append('')

        elif elem.tag.endswith('tbl'):  
            table = doc.tables[[t._element for t in doc.tables].index(elem)]
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    full_text.append(' | '.join(row_text))
                    full_text.append('')
    
    for section in doc.sections:
        for footer in section.footer.paragraphs:
            text = footer.text.strip()
            if text:
                full_text.append(text)
                full_text.append('')
    
    cleaned_text = []
    last_was_empty = False
    for line in full_text:
        if line.strip():
            cleaned_text.append(line)
            last_was_empty = False
        elif not last_was_empty:
            cleaned_text.append('')
            last_was_empty = True
    
    return '\n'.join(cleaned_text).rstrip()


async def extract_text(file_path: str) -> str:
    try:
        file_path = Path(file_path)
        ext = file_path.suffix.lower().lstrip('.')
        
        if ext == "pdf":
            def read_pdf():
                with open(file_path, "rb") as f:
                    reader = PyPDF2.PdfReader(f)
                    return "".join(page.extract_text() or "" for page in reader.pages)
            text = await asyncio.to_thread(read_pdf)
            if not text.strip():
                raise ValueError("Cannot process file: No text extracted")
            return text
        elif ext == "docx":
            text = await extract_docx_text(file_path)
            return text
        elif ext == "xlsx":
            df = pd.read_excel(file_path, engine="openpyxl")
            return df.to_string()
        elif ext == "csv":
            df = pd.read_csv(file_path)
            return df.to_string()
        elif ext == "txt":
            async with aiofiles.open(file_path, "r", encoding="utf-8") as f:
                return await f.read()
        raise ValueError("Unsupported file format")
    
    except Exception as e:
        raise ValueError(f"Error processing file: {e}")
